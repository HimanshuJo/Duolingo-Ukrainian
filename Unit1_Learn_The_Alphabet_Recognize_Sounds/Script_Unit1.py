import os
import glob
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_word_document(input_path, output_path):
    with open(input_path, 'r', encoding='utf-8') as file:
        content = file.read().strip().split('\n\n')

    doc = Document()
    
    # Create a table with 3 columns and apply a table style
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.style.font.name = 'Times New Roman'
    table.style.font.size = Pt(12)  # Adjust the font size as needed
    
    # Apply column header names and style
    for col in table.columns:
        col.width = 3000000  # Adjust the width as needed
    for cell in table.rows[0].cells:
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run()
        run.bold = True
        run.font.size = Pt(14)  # Adjust the font size as needed
        cell.vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        set_cell_shading(cell, "E9E9E9")  # Light gray background color
    table.cell(0, 0).text = 'Ukrainian'
    table.cell(0, 1).text = 'Ukrainian'
    table.cell(0, 2).text = 'English'
    
    for idx, block in enumerate(content):
        lines = block.split('\n')
        if len(lines) == 3:
            row = table.add_row().cells
            for i, cell_text in enumerate(lines):
                row[i].text = cell_text.strip()
                row[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row[i].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                if idx % 2 == 1:  # Alternate row coloring
                    set_cell_shading(row[i], "F0FFFF")
                else:
                    set_cell_shading(row[i], "FFFFFF")
    
    doc.save(output_path)

def set_cell_shading(cell, color_hex):
    shading = cell._tc.get_or_add_tcPr()
    fill = OxmlElement('w:shd')
    fill.set(qn('w:fill'), color_hex)
    shading.append(fill)

def main():
    input_folder = r"C:\CustomProjs\Duolingo-Ukrainian\Unit1_Learn_The_Alphabet_Recognize_Sounds"
    output_folder = r"C:\CustomProjs\Duolingo-Ukrainian\Unit1_Learn_The_Alphabet_Recognize_Sounds"

    os.makedirs(output_folder, exist_ok=True)

    input_files = glob.glob(os.path.join(input_folder, "Part*.txt"))

    for input_file in input_files:
        filename = os.path.splitext(os.path.basename(input_file))[0]
        output_file = os.path.join(output_folder, f"{filename}.docx")

        create_word_document(input_file, output_file)
        print(f"Created: {output_file}")

if __name__ == "__main__":
    main()
